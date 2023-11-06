import docx
from sentence_transformers import SentenceTransformer


class Professional_Standard:
    def __init__(self, path_to_document: str):
        self.doc = docx.Document(path_to_document)
        self.__all_tables = self.doc.tables
        self.__prof_standard_name = self.doc.paragraphs[1].text
        self.__prof_standard_kod = self.__set_prof_standard_kod()
        self.__prof_standard_reg_number = self.__set_prof_standard_reg_number()
        self.__gen_labor_funcs = self.__set_gen_labor_funcs()

        self.__professions, self.__labor_actions,  self.__skills, self.__knowledge = self.__set_other_params()
        self.__knowledge_embeddings = self._get_embeddings(self.__knowledge)
        self.__skills_embeddings = self._get_embeddings(self.__skills)

    def __set_prof_standard_kod(self):
        table = self.__all_tables[2]
        rows = table.rows
        cells = rows[0].cells
        return cells[2].text

    def __set_prof_standard_reg_number(self):
        table = self.__all_tables[1]
        rows = table.rows
        cells = rows[0].cells
        return cells[0].text

    def __set_gen_labor_funcs(self):
        rows = (self.__all_tables[6]).rows
        gen_labor_funcs = {}
        labor_functions = []
        name_gen_labor_func = ''
        for i in range(2, len(rows)):
            cells = rows[i].cells
            if i == 2:
                name_gen_labor_func = cells[1].text
                labor_functions.append(cells[2].text)
                labor_functions.append(cells[3].text)
                continue

            if name_gen_labor_func != cells[1].text:
                gen_labor_funcs[name_gen_labor_func] = labor_functions
                labor_functions = []
                name_gen_labor_func = cells[1].text
                labor_functions.append(cells[2].text)
                labor_functions.append(cells[3].text)
            else:
                labor_functions.append(cells[3].text)
        gen_labor_funcs[name_gen_labor_func] = labor_functions
        return gen_labor_funcs

    def __set_other_params(self):
        i = 9
        tables = self.__all_tables
        professions, labor_actions, skills, knowledge = {}, {}, {}, {}
        for key in self.__gen_labor_funcs.keys():
            table = tables[i]
            for row in table.rows:
                cells = row.cells
                professions[key] = cells[1].text
            i += 5
            for j in range(1, len(self.__gen_labor_funcs[key])):
                _labor_actions, _skills, _knowledge = [], [], []
                values = self.__gen_labor_funcs[key]
                table = self.__all_tables[i]
                for row in table.rows:
                    cells = row.cells
                    if cells[0].text == 'Трудовые действия':
                        if cells[1].text != "-":
                            _labor_actions.append(cells[1].text)
                    if cells[0].text == 'Необходимые умения':
                        if cells[1].text != "-":
                            _skills.append(cells[1].text)
                    if cells[0].text == 'Необходимые знания':
                        if cells[1].text != "-":
                            _knowledge.append(cells[1].text)
                labor_actions[values[j]] = _labor_actions
                skills[values[j]] = _skills
                knowledge[values[j]] = _knowledge
                i += 3
        return professions, labor_actions, skills, knowledge

    def get_prof_standard_name(self):
        return self.__prof_standard_name

    def get_prof_standard_kod(self):
        return self.__prof_standard_kod

    def get_prof_standard_reg_name(self):
        return self.__prof_standard_reg_number

    def get_gen_labor_funcs(self):
        return self.__gen_labor_funcs

    def get_labor_actions(self):
        return self.__labor_actions

    def get_professions(self):
        return self.__professions

    def get_knowledge(self):
        return self.__knowledge

    def get_skills(self):
        return self.__skills

    def get_knowledge_with_embeddings(self):
        return self.__knowledge_embeddings

    def get_skills_with_embeddings(self):
        return self.__skills_embeddings

    @staticmethod
    def _get_embeddings(dict_formulations: dict):
        formulations = []
        for key in dict_formulations.keys():
            for value in dict_formulations[key]:
                if value not in formulations:
                    formulations.append(value)
        model = SentenceTransformer('cointegrated/rubert-tiny2')
        embeddings = model.encode(formulations)
        dict_embeddings = {}
        for formulation, embedding in zip(formulations, embeddings):
            dict_embeddings[formulation] = embedding.tolist()
        return dict_embeddings
